import docx
from docx import Document

def insertCompany(fileName, companyName, role, industry):
    doc = Document(fileName)

    for para in doc.paragraphs:
        if '<COMPANY>' in para.text:
            para.text=para.text.replace('<COMPANY>', companyName)

        if '<ROLE>' in para.text:
            para.text = para.text.replace('<ROLE>', role)

        if '<INDUSTRY>' in para.text:
            para.text = para.text.replace('<INDUSTRY>', industry)

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.name = 'Times New Roman'
            run.font.size = docx.shared.Pt(12)
        
    doc.save(f'created_cover_letters/Andrew Yoo Cover Letter - {companyName}.docx')

def main():
    companyName = input('What is the company name?:\n')
    role = input('What is the role?:\n')
    industry = input('What is the industry?:\n')
    insertCompany("clTemplate.docx", companyName, role, industry)



if __name__ == "__main__":
    main()